from docx import Document
from io import BytesIO

def create(title,summarized_data):
    doc=Document()
    doc.add_heading(title,0)

    for i in list(summarized_data.keys())[:-1]:
        if i!="0":
            doc.add_heading(i,1)
        for j in list(summarized_data[i].keys()):
            if summarized_data[i][j]=="":
                continue
            if j!="0":
                doc.add_heading(j,2)
            doc.add_paragraph(summarized_data[i][j])
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream